import os
import sys
import csv
import json
import logging
import threading
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

# 添加项目根目录到Python路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ExportManager:
    def __init__(self, output_dir: str = None):
        if output_dir is None:
            self.output_dir = Path(__file__).parent.parent.parent / 'exports'
        else:
            self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"导出管理器初始化，输出目录: {self.output_dir}")

    def export_to_excel(
        self,
        records: List[Dict[str, Any]],
        filename: str = None,
        include_charts: bool = False
    ) -> Optional[str]:
        try:
            import pandas as pd
            if not records:
                logger.warning("没有数据可以导出")
                return None

            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"detection_results_{timestamp}.xlsx"

            # 支持完整路径
            if os.path.isabs(filename):
                filepath = Path(filename)
            else:
                filepath = self.output_dir / filename

            df_data = []
            for i, record in enumerate(records, 1):
                df_record = {
                    '序号': i,
                    '检测时间': record.get('timestamp', ''),
                    '垃圾类别(中文)': record.get('class_name_cn', ''),
                    '大类类别': record.get('category_cn', ''),
                    '置信度': record.get('confidence', 0),
                    '位置_X1': '',
                    '位置_Y1': '',
                    '位置_X2': '',
                    '位置_Y2': '',
                    '来源类型': record.get('source_type', 'unknown')
                }
                bbox = record.get('bbox')
                if bbox and len(bbox) == 4:
                    df_record['位置_X1'] = bbox[0]
                    df_record['位置_Y1'] = bbox[1]
                    df_record['位置_X2'] = bbox[2]
                    df_record['位置_Y2'] = bbox[3]
                df_data.append(df_record)

            df = pd.DataFrame(df_data)
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='检测结果', index=False)
                worksheet = writer.sheets['检测结果']
                for idx, col in enumerate(df.columns):
                    max_length = max(df[col].astype(str).apply(len).max(), len(col))
                    worksheet.column_dimensions[chr(65 + idx)].width = min(max_length + 2, 50)
                if include_charts:
                    stats_data = self._generate_stats_data(records)
                    stats_df = pd.DataFrame(stats_data)
                    stats_df.to_excel(writer, sheet_name='统计信息', index=False)
                    stats_sheet = writer.sheets['统计信息']
                    for idx, col in enumerate(stats_df.columns):
                        stats_sheet.column_dimensions[chr(65 + idx)].width = 20

            logger.info(f"Excel文件导出成功: {filepath}")
            return str(filepath)
        except ImportError:
            logger.error("请先安装 pandas 和 openpyxl: pip install pandas openpyxl")
            return None
        except Exception as e:
            logger.error(f"导出Excel失败: {e}")
            return None

    def export_to_csv(
        self,
        records: List[Dict[str, Any]],
        filename: str = None,
        delimiter: str = ','
    ) -> Optional[str]:
        try:
            if not records:
                logger.warning("没有数据可以导出")
                return None

            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"detection_results_{timestamp}.csv"

            if os.path.isabs(filename):
                filepath = Path(filename)
            else:
                filepath = self.output_dir / filename

            with open(filepath, 'w', newline='', encoding='utf-8-sig') as csvfile:
                fieldnames = [
                    '序号', '检测时间', '垃圾类别(中文)', '大类类别',
                    '置信度', '位置_X1', '位置_Y1', '位置_X2', '位置_Y2', '来源类型'
                ]
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=delimiter)
                writer.writeheader()
                for i, record in enumerate(records, 1):
                    csv_record = {
                        '序号': i,
                        '检测时间': record.get('timestamp', ''),
                        '垃圾类别(中文)': record.get('class_name_cn', ''),
                        '大类类别': record.get('category_cn', ''),
                        '置信度': record.get('confidence', 0),
                        '位置_X1': '',
                        '位置_Y1': '',
                        '位置_X2': '',
                        '位置_Y2': '',
                        '来源类型': record.get('source_type', 'unknown')
                    }
                    bbox = record.get('bbox')
                    if bbox and len(bbox) == 4:
                        csv_record['位置_X1'] = bbox[0]
                        csv_record['位置_Y1'] = bbox[1]
                        csv_record['位置_X2'] = bbox[2]
                        csv_record['位置_Y2'] = bbox[3]
                    writer.writerow(csv_record)

            logger.info(f"CSV文件导出成功: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"导出CSV失败: {e}")
            return None

    def export_to_json(
        self,
        records: List[Dict[str, Any]],
        filename: str = None,
        pretty: bool = True
    ) -> Optional[str]:
        try:
            if not records:
                logger.warning("没有数据可以导出")
                return None

            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"detection_results_{timestamp}.json"

            if os.path.isabs(filename):
                filepath = Path(filename)
            else:
                filepath = self.output_dir / filename

            json_data = {
                'export_time': datetime.now().isoformat(),
                'total_records': len(records),
                'records': records
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                if pretty:
                    json.dump(json_data, f, ensure_ascii=False, indent=2)
                else:
                    json.dump(json_data, f, ensure_ascii=False)

            logger.info(f"JSON文件导出成功: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"导出JSON失败: {e}")
            return None

    def export_summary_report(
        self,
        records: List[Dict[str, Any]],
        stats: Dict[str, Any],
        filename: str = None
    ) -> Optional[str]:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml.ns import qn

            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"summary_report_{timestamp}.docx"

            if os.path.isabs(filename):
                filepath = filename
            else:
                filepath = self.output_dir / filename

            doc = Document()
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 标题
            title = doc.add_heading('生活垃圾检测系统 - 检测报告', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 基本信息
            doc.add_heading('基本信息', 1)
            start_time = stats.get('start_time')
            end_time = stats.get('end_time')
            start_str = start_time.strftime('%Y-%m-%d %H:%M:%S') if start_time else 'N/A'
            end_str = end_time.strftime('%Y-%m-%d %H:%M:%S') if end_time else 'N/A'

            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"检测时间: {start_str} - {end_str}")
            doc.add_paragraph(f"总帧数: {stats.get('total_frames', 0)}")
            doc.add_paragraph(f"总检测数: {stats.get('total_detections', 0)}")
            doc.add_paragraph(f"平均置信度: {stats.get('avg_confidence', 0):.2%}")
            doc.add_paragraph(f"检测率: {stats.get('detection_rate', 0):.2%}")

            # 类别统计
            doc.add_heading('类别统计', 1)
            class_counts = stats.get('class_counts', {})
            if class_counts:
                table = doc.add_table(rows=len(class_counts) + 1, cols=3)
                table.style = 'Light Grid Accent 1'
                header_cells = table.rows[0].cells
                header_cells[0].text = '序号'
                header_cells[1].text = '垃圾类别'
                header_cells[2].text = '数量'
                for idx, (class_name, count) in enumerate(sorted(class_counts.items(), key=lambda x: -x[1]), 1):
                    row_cells = table.rows[idx].cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = class_name
                    row_cells[2].text = str(count)

            # 详细检测记录
            doc.add_heading('详细检测记录', 1)
            if records:
                table = doc.add_table(rows=len(records) + 1, cols=5)
                table.style = 'Light Grid Accent 1'
                headers = ['序号', '检测时间', '具体类别', '大类', '置信度']
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                for idx, record in enumerate(records, 1):
                    row_cells = table.rows[idx].cells
                    row_cells[0].text = str(idx)
                    time_str = record.get('timestamp', '')
                    if time_str:
                        try:
                            dt = datetime.fromisoformat(time_str.replace('Z', '+00:00'))
                            time_str = dt.strftime('%H:%M:%S')
                        except:
                            pass
                    row_cells[1].text = time_str
                    row_cells[2].text = record.get('class_name_cn', '')
                    row_cells[3].text = record.get('category_cn', '')
                    confidence = record.get('confidence', 0.0)
                    row_cells[4].text = f"{confidence:.2%}"
            else:
                doc.add_paragraph("无检测记录")

            doc.save(filepath)
            logger.info(f"Word报告导出成功: {filepath}")
            return str(filepath)

        except ImportError:
            logger.error("请先安装 python-docx: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"导出Word报告失败: {e}")
            return None

    def _generate_stats_data(self, records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        from collections import Counter
        category_stats = Counter()
        class_stats = Counter()
        for record in records:
            category = record.get('category_cn', '未知')
            class_name = record.get('class_name_cn', '未知')
            category_stats[category] += 1
            class_stats[class_name] += 1

        stats_data = []
        stats_data.append({'统计类型': '大类类别统计'})
        for category, count in sorted(category_stats.items(), key=lambda x: -x[1]):
            stats_data.append({'统计类型': category, '数量': count})
        stats_data.append({'统计类型': ''})
        stats_data.append({'统计类型': '具体垃圾类别统计'})
        for class_name, count in sorted(class_stats.items(), key=lambda x: -x[1]):
            stats_data.append({'统计类型': class_name, '数量': count})
        return stats_data

    def export_data(self, records: List[Dict[str, Any]], format_type: str, filepath: str = None) -> bool:
        try:
            format_type = format_type.lower()
            if format_type == 'excel':
                result = self.export_to_excel(records, filename=filepath)
                return result is not None
            elif format_type == 'csv':
                result = self.export_to_csv(records, filename=filepath)
                return result is not None
            elif format_type == 'json':
                result = self.export_to_json(records, filename=filepath)
                return result is not None
            elif format_type == 'word':
                # Word 需要 stats，单独调用 export_summary_report
                logger.warning("Word导出请使用 export_summary_report 方法")
                return False
            else:
                logger.error(f"不支持的导出格式: {format_type}")
                return False
        except Exception as e:
            logger.error(f"导出数据失败: {e}")
            return False

    def _generate_simple_stats(self, records: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not records:
            return {
                'total_detections': 0,
                'avg_confidence': 0.0,
                'detection_rate': 0.0,
                'class_counts': {}
            }
        total = len(records)
        avg_conf = sum(r.get('confidence', 0) for r in records) / total
        class_counts = {}
        for record in records:
            class_name = record.get('class_name_cn', '未知')
            class_counts[class_name] = class_counts.get(class_name, 0) + 1
        return {
            'total_detections': total,
            'avg_confidence': avg_conf,
            'detection_rate': 1.0,
            'class_counts': class_counts
        }

    def batch_export(
        self,
        records: List[Dict[str, Any]],
        base_filename: str = None,
        formats: List[str] = None
    ) -> Dict[str, str]:
        if formats is None:
            formats = ['excel', 'csv', 'json']
        if base_filename is None:
            base_filename = f"detection_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        results = {}
        if 'excel' in formats:
            excel_path = self.export_to_excel(records, f"{base_filename}.xlsx")
            if excel_path:
                results['excel'] = excel_path
        if 'csv' in formats:
            csv_path = self.export_to_csv(records, f"{base_filename}.csv")
            if csv_path:
                results['csv'] = csv_path
        if 'json' in formats:
            json_path = self.export_to_json(records, f"{base_filename}.json")
            if json_path:
                results['json'] = json_path
        if 'word' in formats:
            stats = self._generate_simple_stats(records)
            word_path = self.export_summary_report(records, stats, f"{base_filename}.docx")
            if word_path:
                results['word'] = word_path
        return results


_export_manager_instance = None
_export_manager_lock = threading.Lock()


def get_export_manager(output_dir: str = None) -> ExportManager:
    global _export_manager_instance
    with _export_manager_lock:
        if _export_manager_instance is None:
            _export_manager_instance = ExportManager(output_dir)
        return _export_manager_instance
